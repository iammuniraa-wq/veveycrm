"""
VeveyCRM — Documentation Generator
Produces: VeveyCRM_Process_Architecture.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Palette ────────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x15, 0x22, 0x33)   # #152233
BLUE      = RGBColor(0x37, 0x8A, 0xDD)   # #378ADD
AMBER     = RGBColor(0xF6, 0xB2, 0x3C)   # #F6B23C
GREEN     = RGBColor(0x10, 0xB9, 0x81)   # #10B981
RED       = RGBColor(0xEF, 0x44, 0x44)   # #EF4444
TEAL      = RGBColor(0x14, 0xB8, 0xA6)   # #14B8A6
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xF0, 0xF4, 0xF8)
MID_GREY  = RGBColor(0x64, 0x74, 0x8B)


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = "%02X%02X%02X" % (rgb[0], rgb[1], rgb[2])
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def cell_para(cell, text, bold=False, color=None, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    run  = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return para


def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    para.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    para.paragraph_format.space_after  = Pt(6)
    for run in para.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
        run.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)
    return para


def add_body(doc, text, bold=False, color=None, size=10):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return para


def add_bullet(doc, text, level=0, color=None):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent   = Inches(0.25 + level * 0.25)
    para.paragraph_format.space_before  = Pt(1)
    para.paragraph_format.space_after   = Pt(1)
    run = para.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    return para


def flow_table(doc, steps, cols=1):
    """Render a linear flow as a single-column table with arrow connectors."""
    tbl = doc.add_table(rows=0, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    for i, step in enumerate(steps):
        # Step row
        row  = tbl.add_row()
        cell = row.cells[0]
        bg   = step.get("bg", LIGHT_BG)
        set_cell_bg(cell, bg)
        cell.width = Inches(5.5)
        cell_para(cell, step["label"],
                  bold=True, color=step.get("text_color", NAVY),
                  size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        if step.get("sub"):
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(1)
            r2 = p2.add_run(step["sub"])
            r2.font.size = Pt(8.5)
            r2.font.color.rgb = step.get("sub_color", MID_GREY)

        # Arrow row (except after last)
        if i < len(steps) - 1:
            arrow_row  = tbl.add_row()
            arrow_cell = arrow_row.cells[0]
            set_cell_bg(arrow_cell, WHITE)
            cell_para(arrow_cell, "▼", color=BLUE,
                      size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()


def two_col_flow(doc, left_steps, right_steps, left_title="", right_title=""):
    """Side-by-side two-column flow."""
    # Header row
    hdr = doc.add_table(rows=1, cols=2)
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr.style = "Table Grid"
    set_cell_bg(hdr.rows[0].cells[0], NAVY)
    set_cell_bg(hdr.rows[0].cells[1], NAVY)
    cell_para(hdr.rows[0].cells[0], left_title,  bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(hdr.rows[0].cells[1], right_title, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    max_rows = max(len(left_steps), len(right_steps))
    tbl = doc.add_table(rows=max_rows, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    for i in range(max_rows):
        for col_idx, steps in enumerate([left_steps, right_steps]):
            cell = tbl.rows[i].cells[col_idx]
            if i < len(steps):
                step = steps[i]
                set_cell_bg(cell, step.get("bg", LIGHT_BG))
                cell_para(cell, step["label"],
                          bold=step.get("bold", True),
                          color=step.get("text_color", NAVY),
                          size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
                if step.get("sub"):
                    p2 = cell.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r2 = p2.add_run(step["sub"])
                    r2.font.size = Pt(8)
                    r2.font.color.rgb = MID_GREY
            else:
                set_cell_bg(cell, WHITE)
    doc.add_paragraph()


def arch_table(doc, layers):
    """Architecture table: Layer | Components | Notes."""
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = tbl.rows[0].cells
    for cell, txt in zip(hdr_cells, ["Layer", "Components", "Notes"]):
        set_cell_bg(cell, NAVY)
        cell_para(cell, txt, bold=True, color=WHITE, size=10)

    for layer in layers:
        row = tbl.add_row()
        set_cell_bg(row.cells[0], layer.get("bg", LIGHT_BG))
        cell_para(row.cells[0], layer["layer"], bold=True,
                  color=layer.get("label_color", NAVY), size=9.5)
        cell_para(row.cells[1], layer["components"], size=9.5)
        cell_para(row.cells[2], layer.get("notes", ""), size=9, color=MID_GREY, italic=True)
    doc.add_paragraph()


def test_table(doc, tests):
    """Test case table."""
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for cell, txt in zip(tbl.rows[0].cells, ["TC#", "Module", "Scenario", "Expected", "Status"]):
        set_cell_bg(cell, NAVY)
        cell_para(cell, txt, bold=True, color=WHITE, size=9.5)

    for tc in tests:
        row = tbl.add_row()
        status_color = GREEN if tc["status"] == "Pass" else RED if tc["status"] == "Fail" else AMBER
        vals = [tc["id"], tc["module"], tc["scenario"], tc["expected"], tc["status"]]
        bgs  = [WHITE, WHITE, WHITE, WHITE, WHITE]
        for i, (cell, val, bg) in enumerate(zip(row.cells, vals, bgs)):
            set_cell_bg(cell, bg)
            color = status_color if i == 4 else NAVY
            bold  = i == 4
            cell_para(cell, val, size=9, color=color, bold=bold)
    doc.add_paragraph()


def open_points_table(doc, points):
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    for cell, txt in zip(tbl.rows[0].cells, ["#", "Topic", "Description", "Priority"]):
        set_cell_bg(cell, NAVY)
        cell_para(cell, txt, bold=True, color=WHITE, size=9.5)

    for pt in points:
        row = tbl.add_row()
        pri_color = RED if pt["priority"] == "High" else AMBER if pt["priority"] == "Medium" else MID_GREY
        for i, (cell, val) in enumerate(zip(row.cells, [pt["id"], pt["topic"], pt["desc"], pt["priority"]])):
            set_cell_bg(cell, WHITE)
            cell_para(cell, val, size=9,
                      color=pri_color if i == 3 else NAVY,
                      bold=(i == 3))
    doc.add_paragraph()


# ── Build document ──────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("VeveyCRM")
r.font.size  = Pt(32)
r.font.bold  = True
r.font.color.rgb = NAVY

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Process Flows · Architecture · Test Cases · Open Points")
r2.font.size  = Pt(14)
r2.font.color.rgb = BLUE

doc.add_paragraph()

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run(f"Generated: {datetime.date.today().strftime('%d %B %Y')}    ·    Confidential")
r3.font.size  = Pt(10)
r3.font.color.rgb = MID_GREY

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. FULL PROCESS FLOW
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "1. Full End-to-End Process Flow")
add_body(doc,
    "This flow covers the complete lifecycle of a repair job — from the moment a customer "
    "calls in, through service execution, to final invoice and payment. Every step maps to "
    "a VeveyCRM module.",
    size=10)

flow_table(doc, [
    {"label": "Customer Enquiry / Walk-in",
     "sub": "Phone / WhatsApp / walk-in · Lead captured in /leads",
     "bg": RGBColor(0xE0, 0xF2, 0xFE), "text_color": NAVY},
    {"label": "Lead → Account Created",
     "sub": "Account + Contact record created · Type: direct / OEM / end_customer",
     "bg": LIGHT_BG},
    {"label": "Asset Registered",
     "sub": "Motor / Transformer / Pump / Generator logged under account",
     "bg": LIGHT_BG},
    {"label": "Service Case Opened",
     "sub": "Ref: SC-XXXX · Intake date · Equipment label · Complaint noted",
     "bg": RGBColor(0xE0, 0xF2, 0xFE), "text_color": NAVY},
    {"label": "Inspection & Report",
     "sub": "Technician inspects · InspectionReport created · Photos uploaded",
     "bg": LIGHT_BG},
    {"label": "Quotation Raised",
     "sub": "Line items from Pricing Catalogue · PDF generated with Vikas letterhead · Sent to customer",
     "bg": RGBColor(0xFE, 0xF9, 0xC3), "text_color": NAVY},
    {"label": "Quote Approved",
     "sub": "Customer approves (verbal / email / WhatsApp) · Case status → quote_approved",
     "bg": LIGHT_BG},
    {"label": "Work Order Created",
     "sub": "WO-XXXX · Technician assigned · Scheduled date set · Scope of work noted",
     "bg": RGBColor(0xE0, 0xF2, 0xFE), "text_color": NAVY},
    {"label": "Repair Executed",
     "sub": "Technician works on asset · Notes added · Status → in_progress",
     "bg": LIGHT_BG},
    {"label": "Quality Check",
     "sub": "QA stage on service case · Photos + inspection sign-off",
     "bg": LIGHT_BG},
    {"label": "Work Order Marked Complete",
     "sub": "Status → completed · Asset ready for dispatch",
     "bg": RGBColor(0xD1, 0xFA, 0xE5), "text_color": RGBColor(0x06, 0x5F, 0x46)},
    {"label": "Invoice Raised",
     "sub": "INV-XXXX auto-created · Total derived from quote lines · Status → draft",
     "bg": RGBColor(0xD1, 0xFA, 0xE5), "text_color": RGBColor(0x06, 0x5F, 0x46)},
    {"label": "Invoice Sent & Paid",
     "sub": "Status: draft → sent → paid · Payment reconciled",
     "bg": RGBColor(0x14, 0xB8, 0xA6), "text_color": WHITE},
    {"label": "Case Closed",
     "sub": "Service case status → closed · Customer asset returned",
     "bg": NAVY, "text_color": WHITE},
])

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 2. INDIVIDUAL PROCESS FLOWS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "2. Individual Process Flows")

# ── 2.1 Tenant Onboarding ────────────────────────────────────────────────────
add_heading(doc, "2.1  Client (Tenant) Onboarding", level=2)
add_body(doc, "How a new repair shop is added to VeveyCRM as a tenant.")

flow_table(doc, [
    {"label": "Platform Admin logs in to /admin",
     "sub": "admin@veveycrm.com · platform_admins table",
     "bg": RGBColor(0xE0, 0xF2, 0xFE)},
    {"label": "Create New Tenant",
     "sub": "/admin/tenants/new · Company name, slug, accent colour, plan",
     "bg": LIGHT_BG},
    {"label": "Enter Client Admin Email",
     "sub": "One email address — becomes the tenant's first admin user",
     "bg": LIGHT_BG},
    {"label": "Supabase inviteUserByEmail()",
     "sub": "Invite email sent automatically · tenant_users row inserted (role: admin)",
     "bg": RGBColor(0xFE, 0xF9, 0xC3)},
    {"label": "Client Clicks Invite Link",
     "sub": "Sets password · Auth callback fires",
     "bg": LIGHT_BG},
    {"label": "Auth Hook Injects tenant_id into JWT",
     "sub": "Custom Access Token Hook looks up tenant_users · Adds tenant_id claim",
     "bg": RGBColor(0xFE, 0xF9, 0xC3)},
    {"label": "Client Lands in Their Workspace",
     "sub": "RLS scopes all queries to their tenant_id · Completely isolated",
     "bg": RGBColor(0xD1, 0xFA, 0xE5), "text_color": RGBColor(0x06, 0x5F, 0x46)},
    {"label": "Configure Company Info",
     "sub": "Logo, address, GSTIN, partners · Used in PDF letterhead",
     "bg": LIGHT_BG},
])

# ── 2.2 Quotation Flow ───────────────────────────────────────────────────────
add_heading(doc, "2.2  Quotation Flow", level=2)
add_body(doc, "From blank quote to signed-off PDF in the customer's hands.")

flow_table(doc, [
    {"label": "Open /quotations/new",
     "sub": "Select account · Auto-generate ref (QT-YYYY-NNNN)",
     "bg": LIGHT_BG},
    {"label": "Add Line Items",
     "sub": "Description, qty, rate · Pull from Pricing Catalogue · Amount auto-calculated",
     "bg": LIGHT_BG},
    {"label": "Add Terms & Notes",
     "sub": "Free text · Optional valid_until date",
     "bg": LIGHT_BG},
    {"label": "Save Quote",
     "sub": "POST /api/quotes · Inserts quotes + quote_lines + quote_revisions · Returns ID",
     "bg": RGBColor(0xFE, 0xF9, 0xC3)},
    {"label": "Preview PDF",
     "sub": "/quotations/[id]/print · Vikas letterhead + partner logos + line items + totals",
     "bg": LIGHT_BG},
    {"label": "Send to Customer",
     "sub": "Print / email PDF (email integration pending domain registration)",
     "bg": LIGHT_BG},
    {"label": "Customer Approves",
     "sub": "Quote status → approved · Triggers Work Order creation",
     "bg": RGBColor(0xD1, 0xFA, 0xE5), "text_color": RGBColor(0x06, 0x5F, 0x46)},
])

# ── 2.3 Work Order Flow ──────────────────────────────────────────────────────
add_heading(doc, "2.3  Work Order Flow", level=2)

flow_table(doc, [
    {"label": "Work Order Created",
     "sub": "Linked to approved quote or AMC contract · Ref: WO-XXXX",
     "bg": LIGHT_BG},
    {"label": "Technician Assigned",
     "sub": "From technicians pool · Skills matched to job type",
     "bg": LIGHT_BG},
    {"label": "Scheduled",
     "sub": "Status: scheduled · Date/time set · Visible on Dispatch board",
     "bg": RGBColor(0xE0, 0xF2, 0xFE)},
    {"label": "In Progress",
     "sub": "Technician on-site · Status: in_progress · Notes added",
     "bg": RGBColor(0xFE, 0xF9, 0xC3)},
    {"label": "Mark Complete",
     "sub": "PATCH /api/work-orders/[id] → status: completed · Tenant-scoped",
     "bg": RGBColor(0xD1, 0xFA, 0xE5), "text_color": RGBColor(0x06, 0x5F, 0x46)},
    {"label": "Raise Invoice",
     "sub": "POST /api/work-orders/[id]/invoice · INV auto-created · WO → invoiced",
     "bg": RGBColor(0x14, 0xB8, 0xA6), "text_color": WHITE},
])

# ── 2.4 AMC Flow ─────────────────────────────────────────────────────────────
add_heading(doc, "2.4  AMC Contract Flow", level=2)

flow_table(doc, [
    {"label": "AMC Contract Signed",
     "sub": "Contract ref, value, start/end date, account linked",
     "bg": LIGHT_BG},
    {"label": "Periodic Service Scheduled",
     "sub": "Work Order raised against contract (authorized_by.kind = amc)",
     "bg": RGBColor(0xE0, 0xF2, 0xFE)},
    {"label": "Service Executed",
     "sub": "WO flow (see 2.3) — no quotation needed",
     "bg": LIGHT_BG},
    {"label": "Renewal Alert",
     "sub": "AMC page flags contracts expiring ≤60 days in amber",
     "bg": RGBColor(0xFE, 0xF9, 0xC3)},
    {"label": "Renew or Expire",
     "sub": "Status → active / expired / cancelled",
     "bg": LIGHT_BG},
])

# ── 2.5 Invoice Flow ─────────────────────────────────────────────────────────
add_heading(doc, "2.5  Invoice Lifecycle", level=2)

flow_table(doc, [
    {"label": "Draft",
     "sub": "Auto-created when WO is invoiced · Total from quote lines",
     "bg": LIGHT_BG},
    {"label": "Sent",
     "sub": "Invoice PDF shared with customer",
     "bg": RGBColor(0xFE, 0xF9, 0xC3)},
    {"label": "Paid",
     "sub": "Payment confirmed · Status updated",
     "bg": RGBColor(0xD1, 0xFA, 0xE5), "text_color": RGBColor(0x06, 0x5F, 0x46)},
    {"label": "Overdue",
     "sub": "Past due date with no payment · Highlighted in red on /invoices",
     "bg": RGBColor(0xFE, 0xE2, 0xE2), "text_color": RED},
])

# ── 2.6 Dispatch Flow ────────────────────────────────────────────────────────
add_heading(doc, "2.6  Dispatch & Scheduling", level=2)

flow_table(doc, [
    {"label": "Work Orders in scheduled / in_progress status",
     "sub": "Visible on /dispatch board · Ordered by scheduled_for ASC",
     "bg": RGBColor(0xE0, 0xF2, 0xFE)},
    {"label": "Technician availability checked",
     "sub": "technician_leaves table · Skills match",
     "bg": LIGHT_BG},
    {"label": "Assignment confirmed",
     "sub": "WO updated with technician_id + scheduled_for",
     "bg": LIGHT_BG},
    {"label": "Technician dispatched",
     "sub": "Status → in_progress on arrival",
     "bg": RGBColor(0xFE, 0xF9, 0xC3)},
    {"label": "Job completed",
     "sub": "Mark Complete → triggers invoice flow",
     "bg": RGBColor(0xD1, 0xFA, 0xE5), "text_color": RGBColor(0x06, 0x5F, 0x46)},
])

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 3. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "3. System Architecture")

add_heading(doc, "3.1  Technology Stack", level=2)
arch_table(doc, [
    {"layer": "Frontend",        "bg": RGBColor(0xE0, 0xF2, 0xFE), "label_color": BLUE,
     "components": "Next.js 16.2.6 (App Router) · React 19.2.4 · Tailwind CSS 4",
     "notes": "Server components default; client components only for interactivity"},
    {"layer": "Auth",            "bg": LIGHT_BG, "label_color": NAVY,
     "components": "@supabase/ssr · Magic link + email/password · Custom Access Token Hook",
     "notes": "tenant_id injected into JWT on every login"},
    {"layer": "Database",        "bg": LIGHT_BG, "label_color": NAVY,
     "components": "Supabase (PostgreSQL) · 21 tables · Row Level Security on all",
     "notes": "One shared DB; RLS isolates tenants by tenant_id"},
    {"layer": "Storage",         "bg": LIGHT_BG, "label_color": NAVY,
     "components": "Supabase Storage · logos bucket (public)",
     "notes": "Tenant logos, partner logos"},
    {"layer": "API Layer",       "bg": RGBColor(0xE0, 0xF2, 0xFE), "label_color": BLUE,
     "components": "Next.js Route Handlers · /api/* internal · /api/v1/* external REST",
     "notes": "All internal routes gated by requireTenantUser()"},
    {"layer": "Hosting",         "bg": LIGHT_BG, "label_color": NAVY,
     "components": "Vercel · auto-deploy on push to main",
     "notes": "~1 min deploy time"},
    {"layer": "External API",    "bg": LIGHT_BG, "label_color": NAVY,
     "components": "/api/v1/ · Bearer token (VEVEY_API_KEY) · 5 read endpoints",
     "notes": "Currently returns seed data; Supabase wiring pending"},
    {"layer": "MCP",             "bg": LIGHT_BG, "label_color": NAVY,
     "components": "mcp-server/mcp.json · 4 tools for AI assistants",
     "notes": "list_accounts, get_account, list_cases, list_quotations"},
])

add_heading(doc, "3.2  Security Architecture", level=2)
add_body(doc, "Three independent layers — a failure in any one does not expose data.")

arch_table(doc, [
    {"layer": "Layer 1 — JWT Claim",  "bg": RGBColor(0xFE, 0xF9, 0xC3), "label_color": NAVY,
     "components": "Auth Hook → custom_access_token_hook() · Injects tenant_id UUID into JWT",
     "notes": "Runs on every login. Client B's token can never contain Vikas UUID"},
    {"layer": "Layer 2 — RLS",        "bg": RGBColor(0xE0, 0xF2, 0xFE), "label_color": BLUE,
     "components": "PostgreSQL RLS on all 21 tables · using (tenant_id = auth_tenant_id())",
     "notes": "Database engine enforced. Cannot be bypassed by app code"},
    {"layer": "Layer 3 — App Guard",  "bg": RGBColor(0xD1, 0xFA, 0xE5), "label_color": RGBColor(0x06, 0x5F, 0x46),
     "components": "requireTenantUser() · Explicit .eq(tenant_id) on every mutation route",
     "notes": "App-level fallback if RLS ever misconfigured"},
    {"layer": "Admin Routes",         "bg": LIGHT_BG, "label_color": NAVY,
     "components": "isPlatformAdmin() · platform_admins table · Service role key (bypasses RLS)",
     "notes": "Only platform@veveycrm.com level access"},
    {"layer": "External API",         "bg": LIGHT_BG, "label_color": NAVY,
     "components": "VEVEY_API_KEY Bearer token · No fallback if key unset",
     "notes": "Set on Vercel env vars"},
    {"layer": "Dev Backdoor",         "bg": RGBColor(0xD1, 0xFA, 0xE5), "label_color": RGBColor(0x06, 0x5F, 0x46),
     "components": "/api/auth/dev-signin → 404 in production",
     "notes": "Fixed — blocked by NODE_ENV check"},
])

add_heading(doc, "3.3  Data Model Overview", level=2)
add_body(doc, "All domain tables carry tenant_id (FK → tenants). One-liner per entity:")

entities = [
    ("tenants",              "One row per client company. Holds company_info JSONB (letterhead data)."),
    ("tenant_users",         "Links auth users to tenants. role: admin | member."),
    ("platform_admins",      "VeveyCRM platform operators. Bypass RLS via service role."),
    ("accounts",             "Customer companies. type: prospect | oem | direct | end_customer."),
    ("contacts",             "People at an account. phone, email, role."),
    ("assets",               "Equipment. is_loaner flag. loaner_status: available | on_loan."),
    ("service_cases",        "The core repair job record. Has stage timeline, loaner_asset_id, parent_case_id."),
    ("quotes",               "Quotation header. status: draft | sent | approved | rejected."),
    ("quote_lines",          "Line items on a quote. qty × rate = amount."),
    ("quote_revisions",      "Audit log of quote versions."),
    ("contracts",            "AMC contracts. start/end date, value."),
    ("work_orders",          "Field job. authorized_by: {kind: quote|amc, id}. status lifecycle."),
    ("invoices",             "Billing record. status: draft | sent | paid | overdue."),
    ("technicians",          "Field staff. skills, certifications, leave schedule."),
    ("technician_leaves",    "Leave/unavailability periods."),
    ("leads",                "Enquiries. source: oem_referral | amc | direct."),
    ("pricing_items",        "Tenant-specific price catalogue. category, unit, rate."),
    ("text_fragments",       "Reusable text blocks for quotes (terms, disclaimers)."),
    ("visit_logs",           "History of technician visits to accounts."),
    ("activities",           "Activity feed entries (audit trail)."),
    ("case_photos",          "Photos attached to service cases."),
    ("inspection_reports",   "Formal inspection documents on a case."),
]
for name, desc in entities:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    r1 = para.add_run(f"{name}  ")
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = BLUE
    r2 = para.add_run(desc)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = NAVY

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 4. TEST CASES
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "4. Test Cases")
add_body(doc,
    "Critical paths and security scenarios. Status reflects expected behaviour once "
    "Supabase is fully wired and feature flags enabled for all tenants.",
    color=MID_GREY)

test_table(doc, [
    # Onboarding
    {"id": "TC-01", "module": "Onboarding",
     "scenario": "Platform admin creates tenant with valid email",
     "expected": "Tenant row created, invite email sent, tenant_users row inserted",
     "status": "To Test"},
    {"id": "TC-02", "module": "Onboarding",
     "scenario": "Client clicks invite link and sets password",
     "expected": "Auth callback fires, JWT contains tenant_id, lands on dashboard",
     "status": "To Test"},
    {"id": "TC-03", "module": "Onboarding",
     "scenario": "Duplicate slug on tenant creation",
     "expected": "API returns 400 'Slug already taken'",
     "status": "To Test"},

    # Auth & Isolation
    {"id": "TC-04", "module": "Security",
     "scenario": "User from Tenant A tries to read Tenant B's quotes",
     "expected": "Supabase RLS returns 0 rows — no data visible",
     "status": "To Test"},
    {"id": "TC-05", "module": "Security",
     "scenario": "PATCH /api/work-orders/[id] with another tenant's WO id",
     "expected": "404 — .eq(tenant_id) filter returns no match",
     "status": "To Test"},
    {"id": "TC-06", "module": "Security",
     "scenario": "DELETE /api/settings/pricing/[id] for another tenant's item",
     "expected": "404 — tenant_id mismatch",
     "status": "To Test"},
    {"id": "TC-07", "module": "Security",
     "scenario": "GET /api/auth/dev-signin in production",
     "expected": "404 Not found",
     "status": "To Test"},
    {"id": "TC-08", "module": "Security",
     "scenario": "GET /api/v1/accounts with no Authorization header",
     "expected": "401 Unauthorized",
     "status": "To Test"},
    {"id": "TC-09", "module": "Security",
     "scenario": "GET /api/v1/accounts with Authorization: Bearer dev-key",
     "expected": "401 — fallback key removed",
     "status": "To Test"},
    {"id": "TC-10", "module": "Security",
     "scenario": "Unauthenticated user accesses any /app route",
     "expected": "Redirect to /login",
     "status": "To Test"},

    # Quotation
    {"id": "TC-11", "module": "Quotation",
     "scenario": "Create quote with valid account, 3 line items",
     "expected": "Quote + lines + revision row created, ID returned",
     "status": "To Test"},
    {"id": "TC-12", "module": "Quotation",
     "scenario": "Create quote with account_id from another tenant",
     "expected": "404 Account not found",
     "status": "To Test"},
    {"id": "TC-13", "module": "Quotation",
     "scenario": "Create quote with 500 line items",
     "expected": "Only first 200 saved (cap enforced)",
     "status": "To Test"},
    {"id": "TC-14", "module": "Quotation",
     "scenario": "PDF print page loads for saved quote",
     "expected": "Vikas letterhead, partner logos, correct totals",
     "status": "To Test"},

    # Work Order
    {"id": "TC-15", "module": "Work Order",
     "scenario": "Mark Complete on scheduled WO",
     "expected": "Status → completed",
     "status": "To Test"},
    {"id": "TC-16", "module": "Work Order",
     "scenario": "Raise Invoice on completed WO",
     "expected": "Invoice created, WO → invoiced, redirect to /invoices",
     "status": "To Test"},
    {"id": "TC-17", "module": "Work Order",
     "scenario": "Raise Invoice on in_progress WO (not completed)",
     "expected": "400 'Work order must be completed first'",
     "status": "To Test"},

    # Invoices
    {"id": "TC-18", "module": "Invoices",
     "scenario": "/invoices page loads with live data",
     "expected": "KPI strip shows correct ₹ totals per status",
     "status": "To Test"},
    {"id": "TC-19", "module": "Invoices",
     "scenario": "Overdue invoice flagged in red",
     "expected": "Red pill on /invoices table row",
     "status": "To Test"},

    # AMC
    {"id": "TC-20", "module": "AMC",
     "scenario": "Contract expiring in 45 days shown in amber",
     "expected": "Days-left cell rendered amber on /amc",
     "status": "To Test"},
    {"id": "TC-21", "module": "AMC",
     "scenario": "Expired contract shows — in days-left column",
     "expected": "— displayed, not a negative number",
     "status": "To Test"},

    # Settings
    {"id": "TC-22", "module": "Settings",
     "scenario": "Add pricing item via /settings/pricing",
     "expected": "Row inserted with correct tenant_id, appears in list",
     "status": "To Test"},
    {"id": "TC-23", "module": "Settings",
     "scenario": "Delete pricing item",
     "expected": "Row removed, only from own tenant",
     "status": "To Test"},
    {"id": "TC-24", "module": "Settings",
     "scenario": "Add text fragment via /settings/templates",
     "expected": "Row inserted with tenant_id, appears in list",
     "status": "To Test"},

    # Dispatch
    {"id": "TC-25", "module": "Dispatch",
     "scenario": "/dispatch shows only scheduled + in_progress WOs",
     "expected": "Completed/invoiced WOs excluded",
     "status": "To Test"},
    {"id": "TC-26", "module": "Dispatch",
     "scenario": "WO with no technician assigned",
     "expected": "'Unassigned' shown in Technician column",
     "status": "To Test"},

    # PDF
    {"id": "TC-27", "module": "PDF / Print",
     "scenario": "Print page with company_info from Supabase",
     "expected": "Header shows tenant logo, address, GSTIN from DB (not constants fallback)",
     "status": "To Test"},
    {"id": "TC-28", "module": "PDF / Print",
     "scenario": "Partner logo URL is broken/missing",
     "expected": "Text badge fallback shown, no broken image icon",
     "status": "To Test"},
])

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 5. OPEN POINTS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "5. Open Points & Decisions Pending")

open_points_table(doc, [
    {"id": "OP-01", "topic": "Product Name",
     "desc": "'VeveyCRM' is a placeholder — potential trademark conflict with Veeva Systems. Final name decision needed before any public launch or domain registration.",
     "priority": "High"},
    {"id": "OP-02", "topic": "Domain Registration",
     "desc": "No domain owned yet (veveycrm.com or vikaspioneers.com not registered). Blocking email integration — cannot verify sender domain with Resend/SendGrid.",
     "priority": "High"},
    {"id": "OP-03", "topic": "Email Integration",
     "desc": "Resend is the preferred provider. Implementation is straightforward (1 day) but entirely blocked on OP-02. Needed for: quote PDFs, invoice delivery, invite emails (currently generic Supabase template).",
     "priority": "High"},
    {"id": "OP-04", "topic": "Feature Flags — Vikas",
     "desc": "Invoices, Leads, AMC, Dispatch pages are built but the features.invoices / features.leads etc. flags on the Vikas tenant may be false. Enable via /admin/tenants/[id] to show sidebar links.",
     "priority": "High"},
    {"id": "OP-05", "topic": "Seed → Supabase Wiring",
     "desc": "Most list pages (accounts, cases, quotes, work-orders, technicians, reports) still read from seed.ts fixtures. Each needs its data function migrated to a Supabase query in src/lib/data/index.ts.",
     "priority": "High"},
    {"id": "OP-06", "topic": "/api/v1 External REST",
     "desc": "All 5 external API endpoints return seed data. Must be wired to live Supabase with tenant scoping before any integration partner can use them.",
     "priority": "Medium"},
    {"id": "OP-07", "topic": "In-Tenant User Invites",
     "desc": "Only one user per tenant (the email entered at creation). No 'Add team member' flow inside the app. Vikas may need multiple logins for different staff.",
     "priority": "Medium"},
    {"id": "OP-08", "topic": "WhatsApp Integration",
     "desc": "Requires Meta Business Account + WhatsApp Business API approval. Process takes 1-4 weeks. No code written yet. Parked until business decision.",
     "priority": "Medium"},
    {"id": "OP-09", "topic": "Pipeline Page",
     "desc": "Kanban-style sales pipeline. Stub exists at /pipeline. Deliberately deferred — large feature requiring drag-and-drop, stage logic, and probable Supabase real-time.",
     "priority": "Medium"},
    {"id": "OP-10", "topic": "VEVEY_API_KEY on Vercel",
     "desc": "Must be set as a Vercel environment variable. If missing, /api/v1/* returns 401 for all requests. No fallback (intentional after security fix).",
     "priority": "Medium"},
    {"id": "OP-11", "topic": "Supabase Email Templates",
     "desc": "Invite email sent to new tenants uses the default Supabase template. Should be customised in Dashboard → Auth → Email Templates → Invite to include VeveyCRM branding.",
     "priority": "Low"},
    {"id": "OP-12", "topic": "Admin Export — Rate Limiting",
     "desc": "Platform admin can export all tenant data with no rate limit or audit log. Low risk now (1 tenant), but should add logging before scaling.",
     "priority": "Low"},
    {"id": "OP-13", "topic": "Partners Page",
     "desc": "Nav item exists, page is a stub. No spec yet for what this page should show.",
     "priority": "Low"},
    {"id": "OP-14", "topic": "PDF — Send to Customer",
     "desc": "Quote form has Save + PDF Preview. 'Send to customer' button intentionally omitted until email is working (OP-03).",
     "priority": "Low"},
    {"id": "OP-15", "topic": "Mobile — Invoices/Leads/AMC/Dispatch",
     "desc": "Four new pages are desktop-only tables. No responsive layout or .jl-mob CSS applied. If Vikas uses mobile, these pages need a card-view equivalent.",
     "priority": "Low"},
])


# ══════════════════════════════════════════════════════════════════════════════
# FOOTER NOTE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "Document Notes", level=2)
add_body(doc, f"Generated: {datetime.datetime.now().strftime('%d %B %Y, %H:%M')}", color=MID_GREY)
add_body(doc, "Codebase: C:\\Users\\Abdul Nandalpad\\Documents\\veveycrm", color=MID_GREY)
add_body(doc, "Repo: https://github.com/iammuniraa-wq/veveycrm  ·  Live: https://veveycrm-92bd.vercel.app", color=MID_GREY)
add_body(doc, "Stack: Next.js 16.2.6 · React 19.2.4 · Supabase · Vercel", color=MID_GREY)
add_body(doc, "Security layers: JWT Auth Hook (deployed) + RLS on 21 tables + requireTenantUser() app guard", color=MID_GREY)


out_path = r"C:\Users\Abdul Nandalpad\Documents\VeveyCRM_Process_Architecture.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
